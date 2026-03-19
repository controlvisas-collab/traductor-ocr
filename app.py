import streamlit as st
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
from deep_translator import GoogleTranslator
from docx import Document
from io import BytesIO

st.set_page_config(page_title="OCR Traductor Profesional", layout="wide")
st.title("📄 Traductor de Documentos Escaneados")

# Configuración de idiomas
languages = {'Español': 'es', 'Inglés': 'en', 'Francés': 'fr', 'Alemán': 'de', 'Italiano': 'it'}
target_lang = st.sidebar.selectbox("Idioma de destino", list(languages.keys()))

uploaded_file = st.file_uploader("Sube tu PDF escaneado o Imagen", type=["pdf", "png", "jpg", "jpeg"])

if uploaded_file is not None:
    doc_word = Document()
    paginas_imagenes = []
    
    with st.spinner('⏳ Preparando documento...'):
        if uploaded_file.type == "application/pdf":
            # Convertimos PDF a imágenes de alta calidad (300 DPI es ideal para OCR)
            paginas_imagenes = convert_from_bytes(uploaded_file.read(), dpi=300)
        else:
            paginas_imagenes.append(Image.open(uploaded_file))

    translator = GoogleTranslator(source='auto', target=languages[target_lang])

    for i, pagina in enumerate(paginas_imagenes):
        st.markdown(f"### Página {i+1}")
        col1, col2 = st.columns(2)
        
        with col1:
            st.image(pagina, caption="Original", use_container_width=True)
            
        with col2:
            with st.spinner(f'Leyendo y traduciendo página {i+1}...'):
                # Tesseract extrae el texto de la imagen
                texto_extraido = pytesseract.image_to_string(pagina, lang='spa+eng')
                
                if texto_extraido.strip():
                    # Traducimos manteniendo los saltos de párrafo básicos
                    traduccion = translator.translate(texto_extraido)
                    
                    st.text_area(f"Traducción Pág {i+1}", traduccion, height=400)
                    
                    # Añadir al Word con formato de título por página
                    doc_word.add_heading(f'Página {i+1}', level=1)
                    doc_word.add_paragraph(traduccion)
                else:
                    st.warning("No se pudo detectar texto legible en esta página.")

    # Botón de descarga final
    st.divider()
    buffer = BytesIO()
    doc_word.save(buffer)
    buffer.seek(0)
    
    st.download_button(
        label="📥 Descargar Documento Traducido (.docx)",
        data=buffer,
        file_name="traduccion_final.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
